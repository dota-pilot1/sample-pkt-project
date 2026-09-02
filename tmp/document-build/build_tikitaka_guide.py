from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/terecal/pilot-project/sample-pkt-project")
OUT = ROOT / "output/documents/tikitaka-agent-engineering-guide.docx"
BALANCE_IMAGE = Path("/var/folders/4h/z2zcntnn26v1szhfnc6m0t3w0000gn/T/browser-use/assets/25dcbf16-6f07-4b73-9fa9-3a92d3b30bce/e2d31904694b4431.png")
MASCOT_IMAGE = ROOT / "pkt-study-fullstack/public/tikitaka-mascot.png"

# compact_reference_guide preset, with a named Korean-font override.
FONT = "Noto Sans KR"
NAVY = "102A43"
BLUE = "2368B5"
LIGHT_BLUE = "EAF2FB"
RED = "D84A4A"
LIGHT_RED = "FCEEEE"
INK = "182026"
MUTED = "66788A"
LIGHT = "F3F6F9"
LINE = "D9E2EC"
GOLD = "A66A00"


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=None, color=INK, bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=INK, bold=False):
    style.font.name = FONT
    style._element.get_or_add_rPr()
    fonts = style._element.rPr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        style._element.rPr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color=LINE, size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), FONT)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run(run, size=9, color=MUTED)


def new_num_definition(doc, num_id, abstract_id, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "●" if bullet else "%1.")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), FONT)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def add_para(doc, text="", *, size=10.8, bold=False, color=INK, after=6, before=0, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    p.paragraph_format.keep_together = keep
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    apply_num(p, 91)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, size=10.8, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, size=10.8)
    else:
        set_run(p.add_run(text), size=10.8)
    return p


def add_number(doc, text, num_id=92):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    apply_num(p, num_id)
    set_run(p.add_run(text), size=10.8)
    return p


def add_label(doc, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_run(r, size=8.8, color=color, bold=True)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.25
    paragraph_shading(p, fill)
    r1 = p.add_run(f"  {label}  ")
    set_run(r1, size=9.2, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=10.6, color=INK)
    return p


def add_caption(doc, text):
    return add_para(doc, text, size=8.5, color=MUTED, after=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)


def add_section_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    add_bottom_border(p, color=LINE, size="8")
    return p


def add_link_item(doc, label, text, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(f"{label}: "), size=10.2, color=NAVY, bold=True)
    add_hyperlink(p, text, url)
    return p


def configure_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 10.8, INK, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    new_num_definition(doc, 91, 91, bullet=True)
    new_num_definition(doc, 92, 92, bullet=False)
    # A second decimal instance lets independent sequences restart at 1.
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "93")
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), "92")
    num.append(abstract_num_id)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)

    # A third sequence is used by the MES example and also restarts at 1.
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "94")
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), "92")
    num.append(abstract_num_id)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("티키타카 개발 노트  |  Agent Engineering Guide"), size=8.8, color=MUTED, bold=True)
    add_bottom_border(hp, color=LINE, size="6")
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)
    return doc


def build():
    doc = configure_doc()

    # Editorial cover.
    add_para(doc, "FIELD GUIDE", size=9, color=RED, bold=True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "티키타카 개발 노트", size=30, color=NAVY, bold=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    add_para(doc, "공부하고, 만들고, 확인하고, 다시 쓰는 개발 노트", size=15, color=BLUE, bold=True, after=22, align=WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    add_para(doc, "학습 50  :  툴링 50", size=23, color=RED, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "강한 개발자 + 강한 Agent = 더 강한 팀", size=13, color=NAVY, bold=True, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(doc, "핵심", "AI가 코드를 빨리 만들어도 마지막 판단은 사람이 한다. Agent는 대신 일하는 기계가 아니라, 같이 공부하고 확인하는 동료다.", fill=LIGHT, accent=NAVY)
    add_para(doc, "프로젝트 소개 및 사용 가이드", size=10, color=MUTED, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "티키타카 노트 · 데스크톱 학습·개발 노트 앱", size=9, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_label(doc, "PROJECT OVERVIEW")
    add_heading(doc, "0. 프로젝트 한눈에 보기", 1)
    if MASCOT_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        picture = p.add_run().add_picture(str(MASCOT_IMAGE), width=Inches(2.65))
        picture._inline.docPr.set("title", "티키타카 노트 마스코트")
        picture._inline.docPr.set("descr", "Agent를 뜻하는 로봇과 개발 노트를 뜻하는 종이 캐릭터가 함께 서 있는 티키타카 노트 마스코트")
    add_callout(doc, "한 줄 소개", "티키타카 노트는 개발자가 Agent와 함께 공부하고 만든 내용을 한곳에 모아 두고, 다음 작업에서 다시 꺼내 쓰는 데스크톱 노트 앱이다.", fill=LIGHT_BLUE, accent=BLUE)

    add_heading(doc, "누가 쓰면 좋은가", 2)
    for item in [
        "AI 코딩 도구를 쓰지만, 만든 코드를 이해하고 지식도 남기고 싶은 개발자",
        "풀스택이나 MES를 예제로 배우며 작은 기능부터 연습하고 싶은 입문자",
        "프로젝트 규칙, 개발 패턴, 실패 경험을 팀 자산으로 쌓고 싶은 팀",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "무엇이 다른가", 2)
    add_para(doc, "보통 노트는 적고 끝난다. 티키타카 노트는 아래 흐름을 계속 반복한다.")
    add_para(doc, "배운다  →  만든다  →  확인한다  →  적는다  →  다시 쓴다", size=14, color=RED, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_heading(doc, "지금 할 수 있는 일", 2)
    for item in [
        "카테고리와 주제 아래에 본문·하위 문서를 나누어 정리",
        "글, 코드, 표를 편집하고 Agent에게 문서 추가·수정을 요청",
        "MES LOT 예제와 개발 노트를 보며 기능 단위로 학습",
        "SQLite에 로컬 저장하고 앱 안에서 백업·복원",
        "데스크톱 앱 자동 업데이트와 릴리즈 설치 파일 제공",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "시작은 간단하다", 2)
    for step in [
        "최신 릴리즈에서 내 운영체제용 설치 파일을 받는다.",
        "앱을 열고 들어 있는 샘플 노트를 하나 읽는다.",
        "작은 주제를 고른 뒤 직접 적거나 Agent에게 정리를 부탁한다.",
        "코드와 화면을 확인하고, 배운 점을 노트에 남긴다.",
    ]:
        add_number(doc, step)
    add_callout(doc, "기술 구성", "Next.js + Tauri + SQLite로 만든 로컬 우선 데스크톱 앱이다. 인터넷이 없어도 내 노트를 읽고 쓸 수 있고, 데이터는 내 컴퓨터에 저장된다.", fill=LIGHT, accent=NAVY)
    add_heading(doc, "현재 상태", 2)
    for item in [
        "현재 앱 버전: 0.1.58",
        "macOS와 Windows용 설치 패키지 제공",
        "샘플 학습 노트와 기준 데이터를 설치 파일에 함께 포함",
        "다음 단계: Agent 연동, 디자인 노트, 프롬프트·디버깅 지식 모음 강화",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. 왜 이 방식이 필요한가", 1)
    add_para(doc, "AI로 기능을 빨리 만드는 일은 쉬워졌다. 하지만 사람이 결과를 이해하지 못하면 문제가 생긴다. 왜 이렇게 만들었는지 모르고, 틀린 코드가 쌓이고, 장애가 나도 원인을 찾기 어렵다.")
    add_callout(doc, "메시지", "빨리 만드는 것은 시작일 뿐이다. 작은 단위로 만들고, 사람이 확인하고, 배운 것을 적어야 실력이 남는다.", fill=LIGHT_RED, accent=RED)
    add_para(doc, "티키타카 노트는 이 과정을 한곳에서 이어 준다. 공부한 내용과 작업 결과를 남기고, 다음 개발에서 다시 꺼내 쓸 수 있게 만든다.")

    add_heading(doc, "2. 핵심 공식: 학습 50 + 툴링 50", 1)
    if BALANCE_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        picture = p.add_run().add_picture(str(BALANCE_IMAGE), width=Inches(6.35))
        picture._inline.docPr.set("title", "학습 50과 툴링 50의 균형")
        picture._inline.docPr.set("descr", "기본기 학습과 Agent 기반 툴링을 5대5로 균형 있게 반복해 강한 개발자와 강한 팀을 만드는 흐름을 설명하는 인포그래픽")
        add_caption(doc, "프로젝트 핵심 흐름: 기본기와 도구 활용을 5:5로 균형 있게 반복")
    add_para(doc, "학습은 기술 원리, 전체 구조, 코드 읽기, 문제 해결 능력을 키우는 일이다. 툴링은 Agent와 자동화를 써서 공부와 작업 속도를 높이는 일이다. 둘 중 하나만 잘해서는 오래가기 어렵다.")

    add_heading(doc, "3. 티키타카의 다섯 가지 원칙", 1)
    principles = [
        ("1. 기본기부터 쌓는다", "프론트, 백엔드, DB, API를 전부 외울 필요는 없다. 대신 데이터가 어디서 와서 어디로 가는지는 설명할 수 있어야 한다.", "완료한 코드 하나를 골라 입력, 처리, 저장, 출력, 실패 흐름을 내 말로 설명한다."),
        ("2. Agent와 같이 공부한다", "Agent에게 개념을 묻고 코드 검토도 부탁한다. 답만 복사하지 말고, 배운 내용을 짧게 다시 적는다.", "정답과 함께 왜 그런지, 다른 방법은 뭔지, 어떻게 확인할지도 물어본다."),
        ("3. 큰 기능을 한 번에 맡기지 않는다", "목표를 작은 구현 단위로 나누고 각 단위마다 확인한다. 범위가 작을수록 사람도 이해하기 쉽고 Agent도 맥락을 정확히 잡는다.", "한 번의 작업은 한 가지 동작과 하나의 완료 조건을 갖도록 쪼갠다."),
        ("4. 잘 만든 것은 팀의 자산으로 만든다", "코드 규칙, API 방식, DB 설계, 오류 해결법, 실패한 시도를 다시 찾기 쉽게 남긴다.", "결과뿐 아니라 왜 골랐는지와 언제 다시 쓰면 되는지도 적는다."),
        ("5. 쌓은 지식을 다시 Agent에게 준다", "팀의 규칙과 과거 경험을 Agent가 읽을 수 있게 만든다.", "노트를 작게 나누어 두고, 작업을 시작할 때 관련 자료만 Agent에게 준다."),
    ]
    for title, meaning, action in principles:
        add_heading(doc, title, 2)
        meaning_para = add_para(doc, meaning)
        meaning_para.paragraph_format.keep_with_next = True
        add_callout(doc, "바로 해보기", action, fill=LIGHT, accent=BLUE)

    add_heading(doc, "4. 실행 루프: 공부에서 재사용까지", 1)
    add_para(doc, "개발할 때 아래 일곱 단계를 차례로 돈다. 각 단계에서 짧은 기록을 남기면 다음 작업이 쉬워진다.")
    loop_steps = [
        "목표 정의: 이번 작업이 해결할 문제와 완료 조건을 한 문장으로 적는다.",
        "기본기 확인: 관련 개념, 데이터 흐름, 실패 가능성을 먼저 훑는다.",
        "작은 단위 설계: API, 도메인, UI, 테스트를 한꺼번에 묶지 말고 검증 가능한 조각으로 나눈다.",
        "Agent와 구현: 코드와 함께 다른 방법, 장단점, 예상 문제를 질문한다.",
        "사람이 검증: 코드를 읽고 실행하며 테스트, 로그, API 계약, 화면 동작을 확인한다.",
        "노트로 정리: 결정 이유, 검증 증거, 실패한 방법, 재사용 조건을 남긴다.",
        "다시 알려주기: 다음 작업에서 관련 노트를 Agent에게 주고, 새 결과를 또 기록한다.",
    ]
    for step in loop_steps:
        add_number(doc, step, num_id=93)
    add_callout(doc, "루프의 완료 조건", "코드가 동작하는 것만으로 끝내지 않는다. 사람이 설명할 수 있고, 검증 증거가 남아 있으며, 다음 사람이 다시 찾을 수 있어야 완료다.", fill=LIGHT_BLUE, accent=BLUE)

    add_heading(doc, "5. 작은 작업으로 나누는 기준", 1)
    add_para(doc, "작업을 작게 나누면 무엇이 바뀌었는지 알기 쉽고, 문제가 난 곳도 빨리 찾을 수 있다.")
    for item in [
        "완료 조건이 하나인가: '검색된다', '저장된다', '권한이 거부된다'처럼 관찰 가능한 결과가 하나여야 한다.",
        "변경 범위가 설명 가능한가: 무엇을 바꾸고 무엇은 바꾸지 않는지 명확해야 한다.",
        "따로 확인할 수 있는가: 테스트, API 호출, 로그, 화면 중 하나로 바로 확인할 수 있어야 한다.",
        "되돌리기 쉬운가: 실패해도 영향이 제한되고 이전 상태를 복구하기 쉬워야 한다.",
        "지식이 남는가: 구현 뒤에 배운 점과 결정 이유를 짧게라도 기록할 수 있어야 한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. 실전 예시: MES LOT 상태 필터", 1)
    add_para(doc, "MES를 몰라도 괜찮다. 큰 기능을 작은 일로 나누는 흐름만 보면 된다. 'LOT 상태 필터를 전부 만들어줘'라고 한 번에 시키지 말고 아래처럼 나눈다.")
    for item in [
        "계약 먼저: status 쿼리의 허용 값, 기본값, 잘못된 값의 응답을 정의한다.",
        "백엔드 구현: 컨트롤러 파라미터, 조회 조건, 단위 또는 통합 테스트를 추가한다.",
        "API 검증: OpenAPI 문서와 실제 요청 응답에서 파라미터가 노출되고 적용되는지 확인한다.",
        "프론트 구현: 필터 UI, URL 상태, API 호출 값을 연결한다.",
        "화면 검증: 선택값, 새로고침 유지, 빈 결과, 오류 상태를 확인한다.",
        "노트 작성: 계약, 구현 위치, 테스트 명령, 버전 순서와 실패 사례를 남긴다.",
    ]:
        add_number(doc, item, num_id=94)
    add_callout(doc, "Agent에게 물을 질문", "'이 구현에서 조용히 실패할 수 있는 지점은 무엇인가?', 'API 계약과 화면 상태가 어긋났음을 어떻게 검출할 수 있는가?', '테스트를 가장 작은 단위로 나누면 무엇인가?'", fill=LIGHT, accent=NAVY)

    doc.add_page_break()
    add_heading(doc, "7. 다시 쓰기 좋은 노트의 기본 틀", 1)
    add_para(doc, "노트는 길 필요가 없다. 아래 내용을 한 화면에서 찾을 수 있으면 사람도 Agent도 다시 쓰기 쉽다.")
    note_items = [
        ("문제", "무엇이 필요했고 어떤 증상이 있었는가"),
        ("상황", "관련 모듈, 데이터 흐름, 제약 조건, 버전"),
        ("결정", "선택한 방법과 선택 이유"),
        ("안 쓴 방법", "검토했지만 쓰지 않은 방법과 이유"),
        ("바꾼 곳", "핵심 파일, API, 스키마, 코드 패턴"),
        ("검증", "실행한 테스트, 명령, 화면 확인, 로그"),
        ("실패 기록", "틀린 가정, 재현 조건, 해결 단서"),
        ("다시 쓸 때", "언제 다시 적용할 수 있고 언제 적용하면 안 되는가"),
        ("태그", "도메인, 기술, 장애 유형, 관련 기능"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1700, 7660])
    hdr = table.rows[0].cells
    hdr[0].text = "항목"
    hdr[1].text = "반드시 남길 내용"
    for cell in hdr:
        set_cell_fill(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=9.8, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    for label, detail in note_items:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for idx, cell in enumerate(cells):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run(run, size=9.6, color=INK, bold=(idx == 0))
    add_caption(doc, "노트 기본 틀: 결과뿐 아니라 상황, 결정, 확인, 실패, 다시 쓸 조건까지 기록")

    add_heading(doc, "8. Agent에게 줄 자료 묶음", 1)
    add_para(doc, "문서를 전부 주면 Agent도 헷갈린다. 지금 작업에 필요한 자료만 골라서 준다.")
    for item in [
        "프로젝트 규칙: 디렉터리 구조, 의존성 방향, 명명 규칙, 보안 금지사항",
        "전체 구조 그림: 주요 모듈, 데이터 흐름, 외부 연동, 담당 범위",
        "대표 패턴: 인증, 오류 처리, API 응답, DB 변경, 테스트 작성 예시",
        "결정 기록: 기술을 고른 이유와 장단점, 바꾸기 어려운 조건",
        "실패 지식: 반복되는 장애, 재현 절차, 확인할 로그와 지표",
        "현재 상태: 진행 중 작업, 알려진 기술 부채, 다음 검증 포인트",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "운영 원칙", "모든 문서를 한꺼번에 넣지 않는다. 현재 작업과 관련된 최소 문서만 선택하고, 오래된 정보에는 버전과 유효 기간을 표시한다.", fill=LIGHT_RED, accent=RED)

    add_heading(doc, "9. 위험 신호와 대응", 1)
    risks = [
        ("코드는 많은데 설명할 수 없다", "변경을 멈추고 데이터 흐름과 실패 경로를 직접 설명한다."),
        ("한 프롬프트가 여러 계층을 동시에 바꾼다", "API, 도메인, 저장소, UI, 검증을 별도 작업으로 분리한다."),
        ("테스트가 Agent가 만든 코드만 따라간다", "요구사항과 실패 조건에서 독립적으로 테스트 케이스를 다시 만든다."),
        ("노트가 결과 복사본에 그친다", "결정 이유, 대안, 검증 증거, 재사용 조건을 추가한다."),
        ("과거 문서가 현재 코드와 어긋난다", "문서에 기준 버전과 마지막 검증일을 적고 변경 시 함께 갱신한다."),
        ("Agent가 자신 있게 틀린다", "근거 파일과 실행 결과를 요구하고, 불확실한 내용은 가정으로 표시한다."),
    ]
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    set_table_geometry(table2, [3600, 5760])
    headers = table2.rows[0].cells
    headers[0].text = "위험 신호"
    headers[1].text = "즉시 할 대응"
    for cell in headers:
        set_cell_fill(cell, "FCEEEE")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=9.8, color=RED, bold=True)
    set_repeat_table_header(table2.rows[0])
    for risk, response in risks:
        cells = table2.add_row().cells
        cells[0].text = risk
        cells[1].text = response
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run(run, size=9.3)

    doc.add_page_break()
    add_heading(doc, "10. 30일 도입 계획", 1)
    weeks = [
        ("1주차 · 한 기능을 작게", ["진행 중인 기능 하나를 선택한다.", "작업을 3-5개 검증 가능한 단위로 나눈다.", "각 단위에 완료 조건과 검증 방법을 적는다."]),
        ("2주차 · 노트 형식을 고정", ["최소 노트 템플릿을 팀 저장소에 만든다.", "완료한 작업 두 건을 템플릿으로 정리한다.", "실패한 시도 한 건을 반드시 포함한다."]),
        ("3주차 · Agent에 자료 연결", ["프로젝트 규칙과 대표 예시를 한 묶음으로 만든다.", "작업 시작 시 관련 노트를 Agent에게 제공한다.", "답변이 실제 규칙을 반영했는지 확인한다."]),
        ("4주차 · 측정하고 다듬기", ["재작업, 검증 실패, 문서 재사용 사례를 모은다.", "쓸모 없는 항목은 줄이고 빠진 검증 항목은 추가한다.", "다음 달에 자동화할 한 가지를 고른다."]),
    ]
    for week, items in weeks:
        add_heading(doc, week, 2)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "11. 측정 지표", 1)
    add_para(doc, "Agent Engineering의 성과를 코드 생산량만으로 보면 잘못된 최적화를 하기 쉽다. 속도와 함께 이해, 검증, 재사용을 측정한다.")
    metrics = [
        ("이해 가능성", "변경자가 데이터 흐름과 설계 이유를 설명할 수 있는 비율"),
        ("검증 통과율", "첫 구현 뒤 따로 확인했을 때 통과한 작업의 비율"),
        ("작은 배치 시간", "하나의 완료 조건을 구현하고 검증하는 데 걸린 시간"),
        ("지식 재사용률", "새 작업에서 기존 노트나 패턴을 실제로 참조한 비율"),
        ("결함 유출", "검증 단계에서 잡히지 않고 다음 환경으로 넘어간 결함 수"),
        ("최신 문서 비율", "현재 코드와 내용이 맞는 핵심 문서의 비율"),
    ]
    for label, detail in metrics:
        add_bullet(doc, f"{label}: {detail}", bold_prefix=f"{label}:")

    add_heading(doc, "12. 제품 로드맵에 적용하기", 1)
    add_para(doc, "현재 앱은 MES LOT 학습 노트와 Agent 기반 문서 관리를 제공한다. 다음 기능도 메뉴 수를 늘리기보다 '배우고 다시 쓰는 흐름'에 도움이 되는지 보고 고른다.")
    roadmap = [
        ("MCP를 통한 Agent 연동", "필요한 노트를 작업 시점에 검색하고 공급"),
        ("디자인 노트 관리", "UI 결정과 근거, 재사용 가능한 패턴 축적"),
        ("프롬프트 라이브러리", "좋았던 질문 구조와 검증 프롬프트 재사용"),
        ("개발 패턴·코딩 컨벤션", "팀의 기본 구현 방식을 명시"),
        ("디버깅 지식 베이스", "증상, 원인, 로그, 해결 순서를 검색 가능하게 저장"),
        ("기술 부채 관리", "임시 결정과 갚아야 할 조건을 추적"),
        ("팀 지식 직접 참조", "Agent가 현재 규칙과 과거 결정을 근거로 답하도록 연결"),
    ]
    for name, purpose in roadmap:
        add_bullet(doc, f"{name}: {purpose}", bold_prefix=f"{name}:")

    add_heading(doc, "13. Git 및 릴리즈 링크", 1)
    add_para(doc, "아래 주소에서 소스 코드, 설치 파일, 변경 기록, 이슈를 볼 수 있다. 데스크톱 앱의 배포와 업데이트는 pkt-study-fullstack 저장소를 기준으로 한다.")
    add_link_item(doc, "Tauri 앱 소스", "dota-pilot1/pkt-study-fullstack", "https://github.com/dota-pilot1/pkt-study-fullstack")
    add_link_item(doc, "최신 릴리즈", "pkt-study-fullstack/releases/latest", "https://github.com/dota-pilot1/pkt-study-fullstack/releases/latest")
    add_link_item(doc, "전체 릴리즈", "pkt-study-fullstack/releases", "https://github.com/dota-pilot1/pkt-study-fullstack/releases")
    add_link_item(doc, "이슈 및 개선 제안", "pkt-study-fullstack/issues", "https://github.com/dota-pilot1/pkt-study-fullstack/issues")
    add_link_item(doc, "상위 샘플 프로젝트", "dota-pilot1/sample-pkt-project", "https://github.com/dota-pilot1/sample-pkt-project")
    add_callout(doc, "설치 파일 받기", "'최신 릴리즈'를 열고 내 운영체제에 맞는 설치 파일을 받는다. 버전과 게시일도 여기에서 확인할 수 있다.", fill=LIGHT_BLUE, accent=BLUE)

    add_section_rule(doc)
    add_para(doc, "학습 50 + 툴링 50 = Agent Engineering", size=17, color=NAVY, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "사람이 공부하고, Agent와 만들고, 사람이 확인하고, 결과를 적는다. 그리고 다음 작업에서 다시 쓴다. 이 단순한 반복이 티키타카 노트의 핵심이다.", size=11.2, color=INK, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Core metadata and language.
    props = doc.core_properties
    props.title = "티키타카 개발 노트 - Agent Engineering 실전 가이드"
    props.subject = "학습 50과 툴링 50을 결합한 개발 학습 및 지식 자산화 가이드"
    props.author = "티키타카 노트 프로젝트"
    props.keywords = "Agent Engineering, 티키타카 개발 노트, 학습, 툴링, 지식 자산화, MES"
    settings = doc.settings.element
    lang = settings.find(qn("w:themeFontLang"))
    if lang is None:
        lang = OxmlElement("w:themeFontLang")
        settings.append(lang)
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
