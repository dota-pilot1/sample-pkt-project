from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

doc = Document()
for name in ["Noto Sans KR", "Arial Unicode MS", "AppleGothic", "Apple SD Gothic Neo", "AppleMyungjo"]:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}: 티키타카 개발 노트 학습 50 툴링 50 Agent Engineering")
    r.font.name = name
    r.font.size = Pt(18)
    r_pr = r._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)
    fonts.set(qn("w:hint"), "eastAsia")
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")
    lang.set(qn("w:bidi"), "ko-KR")
    r_pr.append(lang)

doc.save("/Users/terecal/pilot-project/sample-pkt-project/tmp/document-build/font-test.docx")
